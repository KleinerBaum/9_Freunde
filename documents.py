"""DocumentAgent für KI-gestützte Berichte via OpenAI Responses API."""

from __future__ import annotations

import random
import time
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from openai import APITimeoutError, OpenAI, OpenAIError, RateLimitError

from config import OpenAIConfig, get_app_config


class DocumentGenerationError(RuntimeError):
    """Domänenspezifischer Fehler bei Dokumentenerstellung."""


class DocumentAgent:
    """Erstellt Berichte auf Basis von Notizen und Kinddaten."""

    def __init__(self) -> None:
        app_config = get_app_config()
        self.openai_config: OpenAIConfig = app_config.openai
        self.client = self._build_client()
        self.logo_path = Path(__file__).resolve().parent / "images" / "logo.png"

    def _build_client(self) -> OpenAI | None:
        if not self.openai_config.api_key:
            return None

        kwargs: dict[str, Any] = {
            "api_key": self.openai_config.api_key,
            "timeout": self.openai_config.timeout_seconds,
        }
        if self.openai_config.base_url:
            kwargs["base_url"] = self.openai_config.base_url
        return OpenAI(**kwargs)

    @property
    def selected_model(self) -> str:
        if self.openai_config.precision_mode == "precise":
            return self.openai_config.model_precise
        return self.openai_config.model_fast

    def _build_tools(self) -> list[dict[str, Any]]:
        tools: list[dict[str, Any]] = []
        if self.openai_config.vector_store_id:
            tools.append(
                {
                    "type": "file_search",
                    "vector_store_ids": [self.openai_config.vector_store_id],
                }
            )
        if self.openai_config.enable_web_search:
            tools.append({"type": "web_search_preview"})
        return tools

    def _response_format(self) -> dict[str, Any]:
        return {
            "type": "json_schema",
            "name": "parent_report",
            "schema": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "body": {"type": "string"},
                },
                "required": ["title", "body"],
                "additionalProperties": False,
            },
            "strict": True,
        }

    def _generate_with_retry(self, prompt: str) -> dict[str, str]:
        if not self.client:
            raise DocumentGenerationError(
                "OpenAI API-Schlüssel fehlt. Bitte [openai].api_key in secrets.toml "
                "oder OPENAI_API_KEY setzen.\n"
                "OpenAI API key is missing. Please set [openai].api_key in secrets.toml "
                "or OPENAI_API_KEY."
            )

        tools = self._build_tools()
        last_error: Exception | None = None

        for attempt in range(self.openai_config.max_retries + 1):
            try:
                response = self.client.responses.create(
                    model=self.selected_model,
                    input=[
                        {
                            "role": "system",
                            "content": "Du bist eine professionelle pädagogische Assistenz. "
                            "Schreibe warmherzige, klare Elternkommunikation.",
                        },
                        {"role": "user", "content": prompt},
                    ],
                    reasoning={"effort": self.openai_config.reasoning_effort},
                    tools=tools,
                    text={"format": self._response_format()},
                )

                payload = response.output_parsed
                if not isinstance(payload, dict):
                    raise DocumentGenerationError(
                        "Die KI-Antwort konnte nicht strukturiert verarbeitet werden."
                    )

                title = str(payload.get("title", "")).strip()
                body = str(payload.get("body", "")).strip()
                if not title or not body:
                    raise DocumentGenerationError(
                        "Die KI-Antwort ist unvollständig. Bitte erneut versuchen."
                    )
                return {"title": title, "body": body}
            except (APITimeoutError, RateLimitError) as exc:
                last_error = exc
            except OpenAIError as exc:
                last_error = exc
                break

            if attempt < self.openai_config.max_retries:
                delay_seconds = min(6.0, (2**attempt) + random.uniform(0.0, 0.4))
                time.sleep(delay_seconds)

        raise DocumentGenerationError(
            "Die Dokumentenerstellung mit OpenAI ist fehlgeschlagen. Bitte später erneut "
            "versuchen oder Konfiguration prüfen.\n"
            "OpenAI document generation failed. Please retry later or verify configuration."
        ) from last_error

    @staticmethod
    def _normalized_language(language: str) -> str:
        return "en" if language.strip().lower() == "en" else "de"

    def generate_document(
        self,
        child_data: dict[str, Any],
        notes: str,
        language: str = "de",
        is_draft: bool = False,
    ) -> tuple[bytes, str]:
        """Generiert einen Dokumenttext mit OpenAI und erstellt ein Word-Dokument."""
        selected_language = self._normalized_language(language)
        child_name = str(child_data.get("name", "Ihr Kind")).strip() or "Ihr Kind"
        language_instruction = (
            "Schreibe den Bericht vollständig auf Deutsch."
            if selected_language == "de"
            else "Write the report fully in English."
        )
        prompt = (
            f"Erstelle einen Elternbericht für {child_name}.\n"
            f"Notizen der Betreuungsperson:\n{notes.strip()}\n\n"
            f"{language_instruction}\n"
            "Nutze einen positiven, klaren Stil. Gib nur valides JSON gemäß Schema zurück."
        )

        result = self._generate_with_retry(prompt)

        doc = Document()
        if self.logo_path.exists():
            doc.add_picture(str(self.logo_path), width=Inches(1.8))
        if is_draft:
            doc.add_paragraph("ENTWURF / DRAFT")
        doc.add_heading(result["title"], level=1)
        today_str = datetime.now().strftime("%d.%m.%Y")
        if selected_language == "de":
            doc.add_paragraph(f"Datum: {today_str}")
        else:
            doc.add_paragraph(f"Date: {today_str}")
        doc.add_paragraph("")
        doc.add_paragraph(result["body"])

        output = BytesIO()
        doc.save(output)
        doc_bytes = output.getvalue()

        safe_name = child_name.replace(" ", "_")
        date_stamp = datetime.now().strftime("%Y%m%d")
        file_prefix = "Bericht" if selected_language == "de" else "Report"
        draft_suffix = "_Entwurf" if is_draft else ""
        file_name = f"{file_prefix}_{safe_name}_{date_stamp}{draft_suffix}.docx"
        return doc_bytes, file_name

    def _add_logo_and_generation_date(self, doc: Document) -> None:
        if self.logo_path.exists():
            doc.add_picture(str(self.logo_path), width=Inches(1.8))
        today_str = datetime.now().strftime("%d.%m.%Y")
        doc.add_paragraph(f"Erstellt am / Generated on: {today_str}")
        doc.add_paragraph("")

    @staticmethod
    def _safe_child_name(child_data: dict[str, Any]) -> str:
        return str(child_data.get("name", "Kind")).strip() or "Kind"

    def generate_care_contract(
        self,
        child_data: dict[str, Any],
        language: str = "de",
        is_draft: bool = False,
    ) -> tuple[bytes, str]:
        """Erstellt einen Betreuungsvertrag auf Basis der Stammdaten."""
        selected_language = self._normalized_language(language)
        child_name = self._safe_child_name(child_data)
        parent_email = str(child_data.get("parent_email", "")).strip() or "—"
        birthdate = str(child_data.get("birthdate", "")).strip() or "—"
        start_date = str(child_data.get("start_date", "")).strip() or "—"
        group = str(child_data.get("group", "")).strip() or "—"
        allergies = str(child_data.get("allergies", "")).strip() or "Keine / None"

        doc = Document()
        self._add_logo_and_generation_date(doc)
        if is_draft:
            doc.add_paragraph("ENTWURF / DRAFT")

        if selected_language == "de":
            doc.add_heading("Betreuungsvertrag", level=1)
            doc.add_paragraph(
                "Zwischen der Großtagespflege 9 Freunde und den Sorgeberechtigten "
                "wird folgender Betreuungsvertrag geschlossen."
            )
            doc.add_heading("1. Vertragsdaten", level=2)
            doc.add_paragraph(f"Kind: {child_name}")
            doc.add_paragraph(f"Elternkontakt: {parent_email}")
            doc.add_paragraph(f"Geburtsdatum: {birthdate}")
            doc.add_paragraph(f"Betreuungsbeginn: {start_date}")
            doc.add_paragraph(f"Gruppe: {group}")
            doc.add_paragraph(f"Allergien: {allergies}")
            doc.add_heading("2. Leistungsumfang", level=2)
            doc.add_paragraph(
                "Die Einrichtung übernimmt die regelmäßige Betreuung, Förderung und "
                "Verpflegung im vereinbarten Betreuungsrahmen."
            )
            doc.add_heading("3. Hinweise", level=2)
        else:
            doc.add_heading("Childcare Contract", level=1)
            doc.add_paragraph(
                "Between Großtagespflege 9 Freunde and the legal guardians the "
                "following childcare contract is concluded."
            )
            doc.add_heading("1. Contract details", level=2)
            doc.add_paragraph(f"Child: {child_name}")
            doc.add_paragraph(f"Parent contact: {parent_email}")
            doc.add_paragraph(f"Birthdate: {birthdate}")
            doc.add_paragraph(f"Start date: {start_date}")
            doc.add_paragraph(f"Group: {group}")
            doc.add_paragraph(f"Allergies: {allergies}")
            doc.add_heading("2. Scope of care", level=2)
            doc.add_paragraph(
                "The daycare provides regular care, educational support and meals "
                "within the agreed scope."
            )
            doc.add_heading("3. Notes", level=2)

        notes_parent = str(child_data.get("notes_parent_visible", "")).strip()
        if selected_language == "de":
            doc.add_paragraph(notes_parent or "Keine zusätzlichen Hinweise.")
        else:
            doc.add_paragraph(notes_parent or "No extra notes.")

        doc.add_paragraph("")
        if selected_language == "de":
            doc.add_paragraph(
                "Ort, Datum: _____________________    Unterschrift Eltern: "
                "_____________________"
            )
            doc.add_paragraph(
                "Ort, Datum: _____________________    Unterschrift Tagespflege: "
                "_____________________"
            )
        else:
            doc.add_paragraph(
                "Place, Date: _____________________    Signature parent: "
                "_____________________"
            )
            doc.add_paragraph(
                "Place, Date: _____________________    Signature daycare: "
                "_____________________"
            )

        output = BytesIO()
        doc.save(output)
        date_stamp = datetime.now().strftime("%Y%m%d")
        file_prefix = "Betreuungsvertrag" if selected_language == "de" else "Contract"
        draft_suffix = "_Entwurf" if is_draft else ""
        file_name = f"{file_prefix}_{child_name.replace(' ', '_')}_{date_stamp}{draft_suffix}.docx"
        return output.getvalue(), file_name

    def generate_food_allowance_invoice(
        self,
        child_data: dict[str, Any],
        period_start: date,
        period_end: date,
        monthly_amount_eur: float,
    ) -> tuple[bytes, str]:
        """Erstellt eine Abrechnung der Lebensmittelpauschale für einen Zeitraum."""
        if period_end < period_start:
            raise DocumentGenerationError(
                "Das Enddatum muss nach dem Startdatum liegen. / End date must be after start date."
            )

        child_name = self._safe_child_name(child_data)
        parent_email = str(child_data.get("parent_email", "")).strip() or "—"
        days_in_period = (period_end - period_start).days + 1
        daily_amount_eur = monthly_amount_eur / 30.0
        total_amount_eur = round(days_in_period * daily_amount_eur, 2)

        doc = Document()
        self._add_logo_and_generation_date(doc)
        doc.add_heading(
            "Abrechnung Lebensmittelpauschale / Food allowance invoice", level=1
        )
        doc.add_paragraph(f"Kind / Child: {child_name}")
        doc.add_paragraph(f"Elternkontakt / Parent contact: {parent_email}")
        doc.add_paragraph(
            "Abrechnungszeitraum / Billing period: "
            f"{period_start.strftime('%d.%m.%Y')} - {period_end.strftime('%d.%m.%Y')}"
        )

        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List"
        table.cell(0, 0).text = "Pauschale pro Monat / Monthly allowance"
        table.cell(0, 1).text = f"{monthly_amount_eur:.2f} €"
        table.cell(1, 0).text = "Tagessatz (Monat/30) / Daily rate (month/30)"
        table.cell(1, 1).text = f"{daily_amount_eur:.2f} €"
        table.cell(2, 0).text = "Anzahl Tage / Number of days"
        table.cell(2, 1).text = str(days_in_period)
        table.cell(3, 0).text = "Gesamtbetrag / Total amount"
        table.cell(3, 1).text = f"{total_amount_eur:.2f} €"

        doc.add_paragraph("")
        doc.add_paragraph(
            "Bitte überweisen Sie den Gesamtbetrag bis zum 10. des Folgemonats. / "
            "Please transfer the total amount by the 10th of the following month."
        )

        output = BytesIO()
        doc.save(output)
        date_stamp = datetime.now().strftime("%Y%m%d")
        file_name = (
            f"Lebensmittelpauschale_{child_name.replace(' ', '_')}_{date_stamp}.docx"
        )
        return output.getvalue(), file_name
