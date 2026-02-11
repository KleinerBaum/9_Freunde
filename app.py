## app.py

from __future__ import annotations

import time
import json
import base64
from io import BytesIO
from pathlib import Path
from typing import Any

from datetime import date

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from googleapiclient.errors import HttpError
from auth import AuthAgent
from stammdaten import StammdatenManager
from documents import DocumentAgent, DocumentGenerationError
from photo import MediaPageContext, PhotoAgent, render_media_page
from storage import DriveAgent
from config import get_app_config, validate_config_or_stop
from services.calendar_service import (
    CalendarServiceError,
    _get_calendar_client,
    _get_calendar_id,
    add_event,
    list_events,
)
from services.drive_service import DriveServiceError, ensure_child_photo_folder
from services.google_clients import get_drive_client
from services.content_repo import ContentRepository, ContentRepositoryError
from services.registration_form_service import (
    RegistrationPayload,
    extract_acroform_fields,
    parse_registration_payload,
)
from services.sheets_repo import (
    SheetsRepositoryError,
    map_schema_v1_payload_to_tab_records,
)
from services.sheets_service import SheetsServiceError, read_sheet_values
from services.photos_service import get_download_bytes

# Streamlit page configuration
LOGO_PATH = Path(__file__).resolve().parent / "images" / "logo.png"
HEART_PATH = Path(__file__).resolve().parent / "images" / "Herz.png"
BACKGROUND_PATH = Path(__file__).resolve().parent / "images" / "Hintergrund.png"
st.set_page_config(
    page_title="9 Freunde App",
    page_icon=str(LOGO_PATH) if LOGO_PATH.exists() else "🤱",
    layout="wide",
)


def _inject_background_image() -> None:
    """Setzt ein dezentes globales Hintergrundbild für die gesamte Streamlit-App."""
    if not BACKGROUND_PATH.exists():
        return

    encoded_background = base64.b64encode(BACKGROUND_PATH.read_bytes()).decode("utf-8")
    st.markdown(
        f"""
        <style>
            [data-testid="stAppViewContainer"] {{
                background-image: linear-gradient(
                        180deg,
                        rgba(255, 255, 255, 0.78) 0%,
                        rgba(255, 255, 255, 0.9) 100%
                    ),
                    url("data:image/png;base64,{encoded_background}");
                background-size: cover;
                background-position: center;
                background-repeat: no-repeat;
                background-attachment: fixed;
            }}

            [data-testid="stHeader"] {{
                background: transparent;
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _trigger_rerun() -> None:
    """Kompatibler Rerun für verschiedene Streamlit-Versionen."""
    rerun_fn = getattr(st, "rerun", None)
    if callable(rerun_fn):
        rerun_fn()
        return

    experimental_rerun_fn = getattr(st, "experimental_rerun", None)
    if callable(experimental_rerun_fn):
        experimental_rerun_fn()


def _parse_optional_iso_date(value: str) -> date | None:
    normalized = value.strip()
    if not normalized:
        return None
    try:
        return date.fromisoformat(normalized)
    except ValueError:
        return None


def _optional_date_to_iso(value: date | None) -> str:
    if value is None:
        return ""
    return value.isoformat()


def _extract_docx_preview_text(
    doc_bytes: bytes,
    *,
    max_paragraphs: int = 12,
) -> str:
    """Erzeugt eine kurze Vorschau aus einem DOCX-Dokument."""
    parsed_document = Document(BytesIO(doc_bytes))
    preview_lines: list[str] = []

    for paragraph in parsed_document.paragraphs:
        normalized_text = paragraph.text.strip()
        if not normalized_text:
            continue
        preview_lines.append(normalized_text)
        if len(preview_lines) >= max_paragraphs:
            break

    if not preview_lines:
        return "Keine Vorschau verfügbar. / No preview available."

    if len(parsed_document.paragraphs) > max_paragraphs:
        preview_lines.append("…")

    return "\n\n".join(preview_lines)


def _normalize_active_flag(value: str | bool | None) -> bool:
    return str(value or "false").strip().lower() == "true"


PHOTO_STATUS_OPTIONS = ("draft", "published", "archived")
DEFAULT_PARENT_VISIBILITY_STATUS = "draft"
GOOGLE_CALENDAR_EMBED_HTML = (
    '<iframe src="https://calendar.google.com/calendar/embed?height=300&wkst=2&ctz='
    "Europe%2FAmsterdam&showPrint=0&showTz=0&showTitle=0&src="
    "NWRmYjdhNTI0YmEzYzYwZWJlODNmOGUyOGNkNDgzYjcxYTEyOTE0MzMxNWI3MGJkMTA2ZTJkZjUy"
    'ZTU2ZmFkZkBncm91cC5jYWxlbmRhci5nb29nbGUuY29t&color=%237986cb" '
    'style="border-width:0" width="800" height="300" frameborder="0" scrolling="no">'
    "</iframe>"
)


def _normalize_photo_status(value: str | None) -> str:
    normalized = str(value or "").strip().lower()
    if normalized in PHOTO_STATUS_OPTIONS:
        return normalized
    return DEFAULT_PARENT_VISIBILITY_STATUS


def _active_flag_to_string(value: bool) -> str:
    return "true" if value else "false"


def _parse_opt_in_flag(value: str | bool | None) -> bool:
    return str(value or "false").strip().lower() == "true"


def _language_label(value: str) -> str:
    normalized = value.strip().lower()
    if normalized == "de":
        return "Deutsch / German"
    if normalized == "en":
        return "English / Englisch"
    return "-"


def _ui_language() -> str:
    language = str(st.session_state.get("ui_language", "de")).strip().lower()
    return "en" if language == "en" else "de"


def _ui_text(value: str) -> str:
    """Reduziert zweisprachige Labels im Format `DE / EN` auf die aktive UI-Sprache."""
    if " / " not in value:
        return value

    german_text, english_text = value.split(" / ", maxsplit=1)
    return english_text if _ui_language() == "en" else german_text


def _display_or_dash(value: object) -> str:
    normalized = str(value or "").strip()
    return normalized if normalized else "-"


def _folder_status_label(child_record: dict[str, str]) -> str:
    has_photo_folder = bool(str(child_record.get("photo_folder_id", "")).strip())
    has_drive_folder = bool(str(child_record.get("folder_id", "")).strip())
    return "✅ Ready" if has_photo_folder or has_drive_folder else "⚠️ Missing"


def _build_admin_overview_rows(
    children: list[dict[str, str]],
    photo_meta_records: list[dict[str, object]],
) -> list[dict[str, str | int]]:
    photo_count_by_child_id: dict[str, int] = {}
    latest_activity_by_child_id: dict[str, str] = {}

    for record in photo_meta_records:
        child_id = str(record.get("child_id", "")).strip()
        if not child_id:
            continue

        photo_count_by_child_id[child_id] = photo_count_by_child_id.get(child_id, 0) + 1

        uploaded_at = str(record.get("uploaded_at", "")).strip()
        if uploaded_at and uploaded_at > latest_activity_by_child_id.get(child_id, ""):
            latest_activity_by_child_id[child_id] = uploaded_at

    overview_rows: list[dict[str, str | int]] = []
    for child_record in children:
        child_id = str(child_record.get("id", "")).strip()
        overview_rows.append(
            {
                "Name / Name": _display_or_dash(child_record.get("name")),
                "Eltern E-Mail / Parent email": _display_or_dash(
                    child_record.get("parent_email")
                ),
                "Fotos / Photos": photo_count_by_child_id.get(child_id, 0),
                "Letzte Aktivität / Last activity": _display_or_dash(
                    latest_activity_by_child_id.get(child_id)
                ),
                "photo_folder_id": _display_or_dash(
                    child_record.get("photo_folder_id")
                ),
                "folder_id": _display_or_dash(child_record.get("folder_id")),
                "Ordnerstatus / Folder status": _folder_status_label(child_record),
            }
        )

    return overview_rows


def _extract_registration_import_data(
    pdf_bytes: bytes,
) -> tuple[RegistrationPayload, dict[str, Any]]:
    fields = extract_acroform_fields(pdf_bytes)
    payload = parse_registration_payload(fields)
    mapped_records = map_schema_v1_payload_to_tab_records(fields)
    return payload, mapped_records


def _save_registration_import(
    stammdaten_manager: StammdatenManager,
    *,
    mapped_records: dict[str, Any],
    created_by: str,
) -> str:
    child_record = dict(mapped_records.get("children", {}))
    parent_records = list(mapped_records.get("parents", []))
    pickup_records = list(mapped_records.get("pickup_authorizations", []))

    child_name = str(child_record.get("name", "")).strip()
    parent_email = str(child_record.get("parent_email", "")).strip()
    if not child_name or not parent_email:
        raise ValueError(
            "Kindname und Parent-E-Mail fehlen im Import. / Child name and parent email are required."
        )

    child_id = stammdaten_manager.add_child(
        child_name,
        parent_email,
        child_record,
    )

    for parent_record in parent_records:
        email = str(parent_record.get("email", "")).strip()
        if not email:
            continue
        parent_payload = {
            key: value for key, value in parent_record.items() if key != "parent_id"
        }
        stammdaten_manager.upsert_parent_by_email(email, parent_payload)

    for pickup_record in pickup_records:
        pickup_payload = {
            key: value
            for key, value in pickup_record.items()
            if key not in {"pickup_id", "child_id"}
        }
        stammdaten_manager.add_pickup_authorization(
            child_id,
            pickup_payload,
            created_by=created_by,
        )

    return child_id


@st.cache_data(show_spinner=False)
def _get_photo_download_bytes(file_id: str, consent_mode: str) -> bytes:
    original_bytes = drive_agent.download_file(file_id)
    return get_download_bytes(original_bytes, consent_mode)


def _run_google_connection_check() -> list[tuple[str, bool, str]]:
    """Prüft Drive-, Calendar- und Sheets-Verbindung mit lesenden Testaufrufen."""
    checks: list[tuple[str, bool, str]] = []
    app_config = get_app_config()

    try:
        drive_client = get_drive_client()
        drive_client.files().list(
            pageSize=1,
            fields="files(id,name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        checks.append(
            (
                "Google Drive Zugriff / Google Drive access",
                True,
                "Drive-Liste erfolgreich gelesen. / Successfully read drive listing.",
            )
        )
    except Exception as exc:  # pragma: no cover - runtime external dependency
        checks.append(
            (
                "Google Drive Zugriff / Google Drive access",
                False,
                "Drive-Test fehlgeschlagen. Prüfen Sie, ob der Service-Account als "
                "Editor auf dem Zielordner eingetragen ist. "
                f"Fehler: {exc}",
            )
        )

    calendar_id = ""
    if app_config.google is not None and isinstance(app_config.google.calendar_id, str):
        calendar_id = app_config.google.calendar_id.strip()

    if app_config.storage_mode == "google" and not calendar_id:
        checks.append(
            (
                "Google Kalender Zugriff / Google Calendar access",
                False,
                "`gcp.calendar_id` fehlt. Hinterlegen Sie die Kalender-ID unter "
                "`Settings → Secrets → [gcp].calendar_id`. / Missing "
                "`gcp.calendar_id`. Add the calendar ID at "
                "`Settings → Secrets → [gcp].calendar_id`.",
            )
        )
    else:
        try:
            calendar_client = _get_calendar_client()
            calendar_client.events().list(
                calendarId=_get_calendar_id(),
                maxResults=1,
                singleEvents=True,
                orderBy="startTime",
            ).execute()
            checks.append(
                (
                    "Google Kalender Zugriff / Google Calendar access",
                    True,
                    "Kalender-Events erfolgreich gelesen. / Successfully read calendar events.",
                )
            )
        except Exception as exc:  # pragma: no cover - runtime external dependency
            checks.append(
                (
                    "Google Kalender Zugriff / Google Calendar access",
                    False,
                    "Kalender-Test fehlgeschlagen. Prüfen Sie, ob die in `calendar_id` "
                    "konfigurierte Kalender-ID mit dem Service-Account geteilt ist "
                    "(mindestens "
                    '"Änderungen an Terminen vornehmen"/"Make changes to events"). '
                    f"Fehler: {exc}",
                )
            )

    def _quote_sheet_tab_for_a1(tab_name: str) -> str:
        escaped = tab_name.replace("'", "''")
        return f"'{escaped}'"

    if app_config.google is not None:
        sheet_id = app_config.google.stammdaten_sheet_id
        sheet_tab = app_config.google.stammdaten_sheet_tab or "children"
        quoted_tab = _quote_sheet_tab_for_a1(sheet_tab)
        check_range = f"{quoted_tab}!A1:A1"
        max_attempts = 3
        last_error: Exception | None = None

        for attempt in range(1, max_attempts + 1):
            try:
                read_sheet_values(sheet_id=sheet_id, range_a1=check_range)
                checks.append(
                    (
                        "Google Sheets Zugriff / Google Sheets access",
                        True,
                        f"Sheets-Lesezugriff erfolgreich (Range {check_range}). / "
                        "Successfully read from Google Sheets "
                        f"(range {check_range}).",
                    )
                )
                break
            except HttpError as exc:  # pragma: no cover - runtime external dependency
                last_error = exc
                status_code = int(getattr(exc.resp, "status", 0) or 0)
                if status_code in (403, 404):
                    break
            except (
                SheetsServiceError
            ) as exc:  # pragma: no cover - runtime external dependency
                last_error = exc
                break
            except Exception as exc:  # pragma: no cover - runtime external dependency
                last_error = exc

            if attempt < max_attempts:
                time.sleep(2 ** (attempt - 1))

        if last_error is not None:
            status_code_value: int | None = None
            if isinstance(last_error, HttpError):
                status_code_value = int(getattr(last_error.resp, "status", 0) or 0)

            if status_code_value == 403:
                message = (
                    "Sheets-Test fehlgeschlagen (403). Die Tabelle ist vermutlich nicht "
                    "mit dem Service-Account geteilt oder die Berechtigung fehlt. / "
                    "Sheets check failed (403). The sheet is likely not shared with the "
                    "service account or permissions are missing."
                )
            elif status_code_value == 404:
                message = (
                    "Sheets-Test fehlgeschlagen (404). Die konfigurierte "
                    "`stammdaten_sheet_id` scheint falsch zu sein. / "
                    "Sheets check failed (404). The configured "
                    "`stammdaten_sheet_id` seems to be invalid."
                )
            else:
                message = (
                    "Sheets-Test fehlgeschlagen. Allgemeiner Google-Sheets-API-Fehler. / "
                    "Sheets check failed. Generic Google Sheets API error."
                )

            checks.append(
                (
                    "Google Sheets Zugriff / Google Sheets access",
                    False,
                    f"{message} Fehler / Error: {last_error}",
                )
            )

    return checks


def _page_title(page: object, language: str) -> str:
    title_de = str(getattr(page, "title_de", "")).strip()
    title_en = str(getattr(page, "title_en", "")).strip()
    if language == "en":
        return title_en or title_de or "Untitled"
    return title_de or title_en or "Ohne Titel"


def _page_body(page: object, language: str) -> str:
    body_de = str(getattr(page, "body_md_de", "")).strip()
    body_en = str(getattr(page, "body_md_en", "")).strip()
    if language == "en":
        return body_en or body_de
    return body_de or body_en


def _normalize_sheet_table(rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    """Normalisiert Google-Sheet-Rohdaten in Header + Datenzeilen."""
    if not rows or not rows[0]:
        return [], []

    header = [str(column).strip() for column in rows[0]]
    data_rows = rows[1:]
    normalized_rows: list[list[str]] = []
    for row in data_rows:
        padded_row = row + [""] * max(0, len(header) - len(row))
        normalized_rows.append(padded_row[: len(header)])
    return header, normalized_rows


def _build_export_payload(rows: list[list[str]]) -> tuple[bytes, bytes] | None:
    """Erstellt CSV- und JSON-Exportdaten aus Sheet-Zeilen."""
    header, normalized_rows = _normalize_sheet_table(rows)
    if not header:
        return None

    dataframe = pd.DataFrame(normalized_rows, columns=header)
    csv_bytes = dataframe.to_csv(index=False).encode("utf-8")
    records = [dict(zip(header, row, strict=False)) for row in normalized_rows]
    json_bytes = json.dumps(records, ensure_ascii=False, indent=2).encode("utf-8")
    return csv_bytes, json_bytes


def _render_export_backup_section() -> None:
    app_config = get_app_config()
    if app_config.storage_mode != "google" or app_config.google is None:
        st.info(
            "Export/Backup ist nur im Google-Modus verfügbar. / "
            "Export/backup is only available in Google mode."
        )
        return

    export_tabs: list[tuple[str, str]] = [
        ("children", app_config.google.children_tab),
        ("parents", app_config.google.parents_tab),
    ]
    st.markdown("**Export / Backup (CSV + JSON)**")
    for export_key, tab_name in export_tabs:
        tab_range = f"{tab_name}!A:ZZ"
        with st.container(border=True):
            st.write(f"**{export_key}** · Tab: `{tab_name}`")
            try:
                export_rows = read_sheet_values(
                    sheet_id=app_config.google.stammdaten_sheet_id,
                    range_a1=tab_range,
                )
            except SheetsServiceError as exc:
                st.warning(
                    "Tab konnte nicht gelesen werden. Bitte Konfiguration prüfen. / "
                    "Could not read tab. Please verify configuration."
                )
                st.caption(str(exc))
                continue
            except Exception as exc:
                st.warning(
                    "Tab konnte nicht gelesen werden (Berechtigung oder Tabname prüfen). / "
                    "Could not read tab (check permissions or tab name)."
                )
                st.caption(str(exc))
                continue

            payload = _build_export_payload(export_rows)
            if payload is None:
                st.info(
                    "Tab ist leer oder enthält keine Header-Zeile. / "
                    "Tab is empty or does not contain a header row."
                )
                continue

            csv_bytes, json_bytes = payload
            csv_col, json_col = st.columns(2)
            with csv_col:
                st.download_button(
                    label="CSV herunterladen / Download CSV",
                    data=csv_bytes,
                    file_name=f"{export_key}.csv",
                    mime="text/csv",
                    key=f"download_export_csv_{export_key}",
                )
            with json_col:
                st.download_button(
                    label="JSON herunterladen / Download JSON",
                    data=json_bytes,
                    file_name=f"{export_key}.json",
                    mime="application/json",
                    key=f"download_export_json_{export_key}",
                )


def _render_child_edit_form(
    child_record: dict[str, str],
    parent_record: dict[str, str],
    key_prefix: str,
    stammdaten_manager: StammdatenManager,
) -> None:
    with st.form(key=f"{key_prefix}_edit_child_form", border=True):
        left_col, right_col = st.columns(2)
        with left_col:
            edit_name = st.text_input(
                "Name des Kindes / Child name",
                value=child_record.get("name", ""),
                key=f"{key_prefix}_edit_name",
            )
            edit_parent_email = st.text_input(
                "E-Mail Elternteil / Parent email",
                value=child_record.get("parent_email", ""),
                key=f"{key_prefix}_edit_parent_email",
            )
            edit_birthdate = st.date_input(
                "Geburtsdatum / Birthdate",
                value=_parse_optional_iso_date(child_record.get("birthdate", "")),
                format="YYYY-MM-DD",
                key=f"{key_prefix}_edit_birthdate",
            )
            edit_start_date = st.date_input(
                "Startdatum / Start date",
                value=_parse_optional_iso_date(child_record.get("start_date", "")),
                format="YYYY-MM-DD",
                key=f"{key_prefix}_edit_start_date",
            )
        with right_col:
            edit_group = st.text_input(
                "Gruppe / Group",
                value=child_record.get("group", ""),
                key=f"{key_prefix}_edit_group",
            )
            edit_primary_caregiver = st.text_input(
                "Bezugserzieher:in / Primary caregiver",
                value=child_record.get("primary_caregiver", ""),
                key=f"{key_prefix}_edit_primary_caregiver",
            )
            edit_allergies = st.text_input(
                "Allergien / Allergies",
                value=child_record.get("allergies", ""),
                key=f"{key_prefix}_edit_allergies",
            )
            edit_pickup_password = st.text_input(
                "Abhol-Kennwort (optional) / Pickup password",
                value=child_record.get("pickup_password", ""),
                type="password",
                key=f"{key_prefix}_edit_pickup_password",
            )

        parent_col_left, parent_col_right = st.columns(2)
        with parent_col_left:
            edit_parent_name = st.text_input(
                "Elternteil Name / Parent name",
                value=str(parent_record.get("name", "")).strip(),
                key=f"{key_prefix}_edit_parent_name",
            )
            edit_parent_phone = st.text_input(
                "Telefon / Phone",
                value=str(parent_record.get("phone", "")).strip(),
                key=f"{key_prefix}_edit_parent_phone",
            )
            edit_parent_phone2 = st.text_input(
                "Telefon 2 / Phone 2",
                value=str(parent_record.get("phone2", "")).strip(),
                key=f"{key_prefix}_edit_parent_phone2",
            )
            edit_parent_address = st.text_input(
                "Adresse / Address",
                value=str(parent_record.get("address", "")).strip(),
                key=f"{key_prefix}_edit_parent_address",
            )
        with parent_col_right:
            edit_emergency_contact_name = st.text_input(
                "Notfallkontakt Name / Emergency contact name",
                value=str(parent_record.get("emergency_contact_name", "")).strip(),
                key=f"{key_prefix}_edit_emergency_contact_name",
            )
            edit_emergency_contact_phone = st.text_input(
                "Notfallkontakt Telefon / Emergency contact phone",
                value=str(parent_record.get("emergency_contact_phone", "")).strip(),
                key=f"{key_prefix}_edit_emergency_contact_phone",
            )
            language_options = ["de", "en"]
            current_parent_language = (
                str(parent_record.get("preferred_language", "de")).strip().lower()
            )
            if current_parent_language not in language_options:
                current_parent_language = "de"
            edit_preferred_language = st.selectbox(
                "Bevorzugte Sprache / Preferred language",
                options=language_options,
                index=language_options.index(current_parent_language),
                format_func=lambda value: (
                    "Deutsch / German" if value == "de" else "English / Englisch"
                ),
                key=f"{key_prefix}_edit_preferred_language",
            )
            edit_notifications_opt_in = st.checkbox(
                "Benachrichtigungen erhalten / Receive notifications",
                value=_parse_opt_in_flag(
                    parent_record.get("notifications_opt_in", "false")
                ),
                key=f"{key_prefix}_edit_notifications_opt_in",
            )

        notes_col_left, notes_col_right = st.columns(2)
        with notes_col_left:
            edit_notes_parent_visible = st.text_area(
                "Hinweise für Eltern sichtbar / Parent-visible notes",
                value=child_record.get("notes_parent_visible", ""),
                height=110,
                key=f"{key_prefix}_edit_notes_parent_visible",
            )
        with notes_col_right:
            edit_notes_internal = st.text_area(
                "Interne Hinweise (nur Leitung) / Internal notes",
                value=child_record.get("notes_internal", ""),
                height=110,
                key=f"{key_prefix}_edit_notes_internal",
            )

        status_options = ["active", "archived"]
        current_status = str(child_record.get("status", "active")).strip()
        if current_status not in status_options:
            current_status = "active"
        edit_status = st.selectbox(
            "Status",
            options=status_options,
            index=status_options.index(current_status),
            key=f"{key_prefix}_edit_status",
        )
        current_download_consent = (
            str(child_record.get("download_consent", "pixelated")).strip().lower()
        )
        consent_options = ["pixelated", "unpixelated", "denied"]
        if current_download_consent not in consent_options:
            current_download_consent = "pixelated"
        edit_download_consent = st.selectbox(
            "Download-Consent / Download consent",
            options=consent_options,
            index=consent_options.index(current_download_consent),
            format_func=lambda mode: (
                "Downloads verpixelt / Downloads pixelated"
                if mode == "pixelated"
                else (
                    "Downloads unverpixelt / Downloads unpixelated"
                    if mode == "unpixelated"
                    else "Download verweigert / Download denied"
                )
            ),
            key=f"{key_prefix}_edit_download_consent",
        )
        edit_submitted = st.form_submit_button("Änderungen speichern / Save")

    if not edit_submitted:
        return

    if not edit_name.strip() or not edit_parent_email.strip():
        st.error(
            "Bitte Name und Eltern-E-Mail angeben. / "
            "Please provide child name and parent email."
        )
        return

    try:
        stammdaten_manager.update_child(
            child_record.get("id", ""),
            {
                "name": edit_name.strip(),
                "parent_email": edit_parent_email.strip(),
                "birthdate": _optional_date_to_iso(edit_birthdate),
                "start_date": _optional_date_to_iso(edit_start_date),
                "group": edit_group.strip(),
                "primary_caregiver": edit_primary_caregiver.strip(),
                "allergies": edit_allergies.strip(),
                "notes_parent_visible": edit_notes_parent_visible.strip(),
                "notes_internal": edit_notes_internal.strip(),
                "pickup_password": edit_pickup_password.strip(),
                "status": edit_status,
                "download_consent": edit_download_consent,
            },
        )
        stammdaten_manager.upsert_parent_by_email(
            edit_parent_email.strip(),
            {
                "name": edit_parent_name.strip(),
                "phone": edit_parent_phone.strip(),
                "phone2": edit_parent_phone2.strip(),
                "address": edit_parent_address.strip(),
                "preferred_language": edit_preferred_language,
                "emergency_contact_name": edit_emergency_contact_name.strip(),
                "emergency_contact_phone": edit_emergency_contact_phone.strip(),
                "notifications_opt_in": _active_flag_to_string(
                    edit_notifications_opt_in
                ),
            },
        )
        st.success("Kind wurde aktualisiert. / Child record updated.")
        _trigger_rerun()
    except Exception as exc:
        st.error(f"Speichern fehlgeschlagen / Save failed: {exc}")


# Validate required secrets early and fail with clear UI guidance
validate_config_or_stop()
app_config = get_app_config()
_inject_background_image()

if HEART_PATH.exists():
    col_left, col_center, col_right = st.columns([1, 2, 1])
    with col_center:
        st.image(str(HEART_PATH), width=180)

if LOGO_PATH.exists():
    st.image(str(LOGO_PATH), width=180)

# Initialize agents (ensure single instance per session)
if "auth_agent" not in st.session_state:
    st.session_state.auth_agent = AuthAgent()
if "stammdaten_manager" not in st.session_state:
    st.session_state.stammdaten_manager = StammdatenManager()
if "drive_agent" not in st.session_state:
    st.session_state.drive_agent = DriveAgent()
if "doc_agent" not in st.session_state:
    st.session_state.doc_agent = DocumentAgent()
if "photo_agent" not in st.session_state:
    st.session_state.photo_agent = PhotoAgent()

auth = st.session_state.auth_agent
stammdaten_manager = st.session_state.stammdaten_manager
drive_agent = st.session_state.drive_agent
doc_agent = st.session_state.doc_agent
photo_agent = st.session_state.photo_agent
content_repo = ContentRepository()

# Check if user is logged in
if "user" not in st.session_state or st.session_state.get("user") is None:
    # --- Login Screen ---
    st.title("9 Freunde – Login")
    email = st.text_input("E-Mail")
    password = st.text_input("Passwort", type="password")
    if st.button("Anmelden"):
        user_role = auth.login(email, password)
        if user_role:
            # Successful login
            st.session_state.user = email
            st.session_state.role = user_role  # "admin" oder "parent"
            # Wenn Elternrolle: zugehöriges Kind ermitteln
            if user_role == "parent":
                child = stammdaten_manager.get_child_by_parent(email)
                if child:
                    st.session_state.child = child
                else:
                    st.session_state.child = None
            _trigger_rerun()
        else:
            st.error("Login fehlgeschlagen. Bitte E-Mail/Passwort überprüfen.")
else:
    # --- Main Application (Post-Login) ---
    user_role = st.session_state.role
    user_email = st.session_state.user

    # Sidebar menu based on role
    st.sidebar.title("9 Freunde App")
    st.sidebar.radio(
        "Sprache / Language",
        options=("de", "en"),
        format_func=_language_label,
        horizontal=True,
        key="ui_language",
    )
    st.sidebar.write(f"{_ui_text('Angemeldet als / Logged in as')}: `{user_email}`")
    healthcheck_requested = False
    if user_role == "admin":
        admin_menu_labels: dict[str, str] = {
            "dashboard": "Dashboard / Dashboard",
            "master_data": "Stammdaten & Infos / Master data & info",
            "photos": "Fotos & Medien / Photos & media",
            "documents": "Dokumente & Verträge / Documents & contracts",
            "calendar": "Kalender / Calendar",
            "system": "System / Healthchecks",
        }
        menu = st.sidebar.radio(
            _ui_text("Hauptnavigation / Main navigation"),
            options=tuple(admin_menu_labels.keys()),
            format_func=lambda key: _ui_text(admin_menu_labels[key]),
            index=0,
        )
    else:
        parent_menu_labels: dict[str, str] = {
            "child": "Mein Kind / My child",
            "info": "Infos / Info",
            "documents": "Dokumente / Documents",
            "photos": "Fotos / Photos",
            "appointments": "Termine / Appointments",
            "medication": "Medikationen / Medication",
        }
        menu = st.sidebar.radio(
            _ui_text("Menü / Menu"),
            options=tuple(parent_menu_labels.keys()),
            format_func=lambda key: _ui_text(parent_menu_labels[key]),
            index=0,
        )

    if user_role == "admin" and menu == "system":
        healthcheck_requested = st.sidebar.button(
            _ui_text("Google-Verbindung prüfen / Check Google connection")
        )

    # Logout button at bottom of sidebar
    if st.sidebar.button(_ui_text("Abmelden / Logout")):
        # Clear session state and rerun to show login
        st.session_state.clear()
        _trigger_rerun()

    # Content area:
    if user_role == "admin":
        st.header(_ui_text(admin_menu_labels[menu]))
    else:
        st.header(_ui_text(parent_menu_labels[menu]))

    if user_role == "admin":
        admin_view = menu
        if menu == "dashboard":
            admin_view = "Dashboard"
        elif menu == "master_data":
            admin_view = "Stammdaten"
        elif menu == "documents":
            admin_view = st.radio(
                _ui_text("Bereich / Section"),
                ("Dokumente", "Verträge"),
                horizontal=True,
                key="admin_documents_section",
            )
        elif menu == "photos":
            admin_view = "Fotos"
        elif menu == "calendar":
            admin_view = "Kalender"
        elif menu == "system":
            admin_view = "System / Healthchecks"

        # ---- Admin: Dashboard ----
        if admin_view == "Dashboard":
            with st.container(border=True):
                st.subheader("Dashboard / Dashboard")
                try:
                    dashboard_children = stammdaten_manager.get_children()
                    active_children_count = len(
                        [
                            child
                            for child in dashboard_children
                            if str(child.get("status", "active")).strip().lower()
                            == "active"
                        ]
                    )
                except Exception:
                    dashboard_children = []
                    active_children_count = 0

                col_total, col_active, col_inactive = st.columns(3)
                with col_total:
                    st.metric("Kinder gesamt / Total children", len(dashboard_children))
                with col_active:
                    st.metric("Aktiv / Active", active_children_count)
                with col_inactive:
                    st.metric(
                        "Archiviert / Archived",
                        len(dashboard_children) - active_children_count,
                    )

                st.subheader("Admin-Übersicht / Admin overview")
                children: list[dict[str, str]] = []
                children_load_error = False
                try:
                    children = stammdaten_manager.get_children()
                    photo_meta_records = stammdaten_manager.get_photo_meta_records()
                except SheetsRepositoryError as exc:
                    children_load_error = True
                    st.error(
                        "Stammdaten konnten nicht geladen werden. Bitte prüfen Sie, ob der "
                        "Service-Account Zugriff auf die konfigurierte Tabelle hat "
                        "(gcp.stammdaten_sheet_id). / Could not load master data. Please "
                        "verify that the service account has access to the configured "
                        "sheet (gcp.stammdaten_sheet_id)."
                    )
                    st.caption(f"Details / Details: {exc}")
                    photo_meta_records = []
                except Exception:
                    children_load_error = True
                    st.error(
                        "Stammdaten konnten aktuell nicht geladen werden. Bitte später "
                        "erneut versuchen. / Master data could not be loaded right now. "
                        "Please try again later."
                    )
                    photo_meta_records = []

                if not children_load_error and children:
                    st.markdown("**👥 Kinder-Übersicht / Children overview**")
                    overview_df = pd.DataFrame(
                        _build_admin_overview_rows(children, photo_meta_records)
                    )
                    st.dataframe(
                        overview_df,
                        hide_index=True,
                        width="stretch",
                    )
                elif not children_load_error:
                    st.write(
                        "*Noch keine Kinder registriert. / No children registered yet.*"
                    )

                st.info(
                    "Das Dashboard zeigt Kennzahlen und die aktuelle Kinder-Übersicht auf "
                    "einen Blick. / The dashboard shows key metrics and the current children "
                    "overview at a glance."
                )

        # ---- Admin: Stammdaten ----
        if admin_view == "Stammdaten":
            st.subheader("Kinder-Stammdaten verwalten")
            children: list[dict[str, str]] = []
            children_load_error = False
            try:
                children = stammdaten_manager.get_children()
            except SheetsRepositoryError as exc:
                children_load_error = True
                st.error(
                    "Stammdaten konnten nicht geladen werden. Bitte prüfen Sie, ob der "
                    "Service-Account Zugriff auf die konfigurierte Tabelle hat "
                    "(gcp.stammdaten_sheet_id). / Could not load master data. Please "
                    "verify that the service account has access to the configured "
                    "sheet (gcp.stammdaten_sheet_id)."
                )
                st.caption(f"Details / Details: {exc}")
            except Exception:
                children_load_error = True
                st.error(
                    "Stammdaten konnten aktuell nicht geladen werden. Bitte später "
                    "erneut versuchen. / Master data could not be loaded right now. "
                    "Please try again later."
                )

            if not children_load_error and children:
                st.markdown("**👥 Kinder-Übersicht / Children overview**")
                overview_rows: list[dict[str, str]] = []
                for child_record in children:
                    overview_rows.append(
                        {
                            "Name": _display_or_dash(child_record.get("name")),
                            "Parent Email": _display_or_dash(
                                child_record.get("parent_email")
                            ),
                            "Group": _display_or_dash(child_record.get("group")),
                            "Birthdate": _display_or_dash(
                                child_record.get("birthdate")
                            ),
                            "Folder Status": _folder_status_label(child_record),
                        }
                    )
                overview_df = pd.DataFrame(overview_rows)
                st.dataframe(
                    overview_df,
                    hide_index=True,
                    width="stretch",
                )
            elif not children_load_error:
                st.write("*Noch keine Kinder registriert.*")

            # Formular zum Hinzufügen eines neuen Kindes
            with st.expander("Neues Kind anlegen / Add child", expanded=False):
                with st.form(key="new_child_form", border=True):
                    left_col, right_col = st.columns(2)
                    with left_col:
                        name = st.text_input("Name des Kindes / Child name")
                        parent_email = st.text_input("E-Mail Elternteil / Parent email")
                        birthdate = st.date_input(
                            "Geburtsdatum / Birthdate",
                            value=None,
                            format="YYYY-MM-DD",
                        )
                        start_date = st.date_input(
                            "Startdatum / Start date",
                            value=None,
                            format="YYYY-MM-DD",
                        )
                        status = st.selectbox(
                            "Status", options=["active", "archived"], index=0
                        )
                    with right_col:
                        group = st.text_input("Gruppe / Group", value="Igel")
                        primary_caregiver = st.text_input(
                            "Bezugserzieher:in / Primary caregiver"
                        )
                        allergies = st.text_input("Allergien / Allergies")
                        pickup_password = st.text_input(
                            "Abhol-Kennwort (optional) / Pickup password",
                            type="password",
                        )

                    parent_col_left, parent_col_right = st.columns(2)
                    with parent_col_left:
                        parent_name = st.text_input("Elternteil Name / Parent name")
                        parent_phone = st.text_input("Telefon / Phone")
                        parent_phone2 = st.text_input("Telefon 2 / Phone 2")
                        parent_address = st.text_input("Adresse / Address")
                    with parent_col_right:
                        emergency_contact_name = st.text_input(
                            "Notfallkontakt Name / Emergency contact name"
                        )
                        emergency_contact_phone = st.text_input(
                            "Notfallkontakt Telefon / Emergency contact phone"
                        )
                        preferred_language = st.selectbox(
                            "Bevorzugte Sprache / Preferred language",
                            options=["de", "en"],
                            index=0,
                            format_func=lambda value: (
                                "Deutsch / German"
                                if value == "de"
                                else "English / Englisch"
                            ),
                        )
                        notifications_opt_in = st.checkbox(
                            "Benachrichtigungen erhalten / Receive notifications",
                            value=False,
                        )

                    notes_col_left, notes_col_right = st.columns(2)
                    with notes_col_left:
                        notes_parent_visible = st.text_area(
                            "Hinweise für Eltern sichtbar / Parent-visible notes",
                            height=110,
                        )
                    with notes_col_right:
                        notes_internal = st.text_area(
                            "Interne Hinweise (nur Leitung) / Internal notes",
                            height=110,
                        )
                    submitted = st.form_submit_button("Hinzufügen / Add child")
            if submitted:
                if name.strip() == "" or parent_email.strip() == "":
                    st.error(
                        "Bitte Name und Eltern-E-Mail angeben. / Please provide child name and parent email."
                    )
                else:
                    try:
                        stammdaten_manager.add_child(
                            name.strip(),
                            parent_email.strip(),
                            {
                                "birthdate": _optional_date_to_iso(birthdate),
                                "start_date": _optional_date_to_iso(start_date),
                                "group": group.strip(),
                                "primary_caregiver": primary_caregiver.strip(),
                                "allergies": allergies.strip(),
                                "notes_parent_visible": notes_parent_visible.strip(),
                                "notes_internal": notes_internal.strip(),
                                "pickup_password": pickup_password.strip(),
                                "status": status,
                            },
                        )
                        stammdaten_manager.upsert_parent_by_email(
                            parent_email.strip(),
                            {
                                "name": parent_name.strip(),
                                "phone": parent_phone.strip(),
                                "phone2": parent_phone2.strip(),
                                "address": parent_address.strip(),
                                "preferred_language": preferred_language,
                                "emergency_contact_name": emergency_contact_name.strip(),
                                "emergency_contact_phone": emergency_contact_phone.strip(),
                                "notifications_opt_in": _active_flag_to_string(
                                    notifications_opt_in
                                ),
                            },
                        )
                        st.success(
                            f"Kind '{name}' hinzugefügt. / Child '{name}' added."
                        )
                        _trigger_rerun()
                    except Exception as e:
                        st.error(f"Fehler beim Speichern / Save error: {e}")

            with st.expander(
                "Anmeldung importieren (PDF) / Import registration (PDF)",
                expanded=False,
            ):
                template_path = (
                    Path(__file__).resolve().parent
                    / "assets"
                    / "forms"
                    / "9Freunde_Anmeldeformular_v1.pdf"
                )
                if template_path.exists():
                    st.download_button(
                        "Leere Vorlage herunterladen / Download blank template",
                        data=template_path.read_bytes(),
                        file_name=template_path.name,
                        mime="application/pdf",
                        key="registration_template_download",
                    )
                else:
                    st.warning("Vorlage nicht gefunden. / Template file not found.")

                with st.form(key="registration_import_form", border=True):
                    uploaded_registration_pdf = st.file_uploader(
                        "Ausgefülltes Anmeldeformular hochladen / Upload completed registration form",
                        type=["pdf"],
                        key="registration_import_pdf",
                    )
                    import_submitted = st.form_submit_button(
                        "In Stammdaten speichern / Save to master data",
                        type="primary",
                    )

                if import_submitted:
                    if uploaded_registration_pdf is None:
                        st.warning(
                            "Bitte zuerst ein PDF auswählen. / Please select a PDF file first."
                        )
                    else:
                        registration_payload: RegistrationPayload | None = None
                        mapped_registration_records: dict[str, Any] | None = None
                        import_errors: list[str] = []

                        try:
                            (
                                registration_payload,
                                mapped_registration_records,
                            ) = _extract_registration_import_data(
                                uploaded_registration_pdf.getvalue()
                            )
                            import_errors = list(registration_payload.errors)
                        except ValueError as exc:
                            import_errors = [str(exc)]
                        except Exception as exc:
                            import_errors = [
                                "Import konnte nicht verarbeitet werden. / Import could not be processed.",
                                f"Details / Details: {exc}",
                            ]

                        if registration_payload is not None:
                            child_preview = registration_payload.child
                            parent_preview = registration_payload.parents
                            pickup_preview = registration_payload.pickup_authorizations
                            consent_preview = registration_payload.consents

                            st.markdown("**Kind-Zusammenfassung / Child summary**")
                            st.write(
                                {
                                    "Name": _display_or_dash(child_preview.get("name")),
                                    "Geburtsdatum / Birthdate": _display_or_dash(
                                        child_preview.get("birthdate")
                                    ),
                                    "Startdatum / Start date": _display_or_dash(
                                        child_preview.get("start_date")
                                    ),
                                    "Gruppe / Group": _display_or_dash(
                                        child_preview.get("group")
                                    ),
                                    "Parent E-Mail": _display_or_dash(
                                        child_preview.get("parent_email")
                                        or child_preview.get("parent1_email")
                                    ),
                                }
                            )

                            st.markdown("**Eltern-Zusammenfassung / Parent summary**")
                            if parent_preview:
                                st.dataframe(
                                    pd.DataFrame(parent_preview),
                                    hide_index=True,
                                    width="stretch",
                                )
                            else:
                                st.info(
                                    "Keine Elternangaben gefunden. / No parent data found."
                                )

                            st.markdown(
                                "**Abholberechtigte / Pickup authorization list**"
                            )
                            if pickup_preview:
                                st.dataframe(
                                    pd.DataFrame(pickup_preview),
                                    hide_index=True,
                                    width="stretch",
                                )
                            else:
                                st.info(
                                    "Keine Abholberechtigten erkannt. / No pickup authorizations detected."
                                )

                            st.markdown("**Einwilligungen / Consent summary**")
                            st.write(consent_preview)

                        if import_errors:
                            st.error(
                                "Validierungsfehler gefunden. Speichern ist blockiert. / Validation errors found. Saving is blocked."
                            )
                            for error in import_errors:
                                st.write(f"- {error}")
                        elif mapped_registration_records is not None:
                            try:
                                imported_child_id = _save_registration_import(
                                    stammdaten_manager,
                                    mapped_records=mapped_registration_records,
                                    created_by=user_email,
                                )
                                st.session_state["stammdaten_selected_child_ids"] = [
                                    imported_child_id
                                ]
                                if hasattr(st, "toast"):
                                    st.toast(
                                        "Anmeldung erfolgreich importiert. / Registration imported successfully."
                                    )
                                st.success(
                                    "Import erfolgreich gespeichert. / Import saved successfully. "
                                    f"child_id: `{imported_child_id}`"
                                )
                                st.caption(
                                    "Direkt zur Bearbeitung vorausgewählt. / Preselected for direct editing."
                                )
                                _trigger_rerun()
                            except Exception as exc:
                                st.error(
                                    "Fehler beim Speichern des Imports. / Error while saving import."
                                )
                                st.caption(f"Details / Details: {exc}")
            if children:
                st.write("**Kind bearbeiten / Edit child:**")
                selected_ids = set(
                    st.session_state.get("stammdaten_selected_child_ids", [])
                )
                selection_rows: list[dict[str, str | bool]] = []
                for child_record in children:
                    child_id = str(child_record.get("id", "")).strip()
                    selection_rows.append(
                        {
                            "Auswahl / Select": child_id in selected_ids,
                            "Name / Name": _display_or_dash(child_record.get("name")),
                            "Parent Email": _display_or_dash(
                                child_record.get("parent_email")
                            ),
                            "Group": _display_or_dash(child_record.get("group")),
                            "Birthdate": _display_or_dash(
                                child_record.get("birthdate")
                            ),
                            "Folder Status": _folder_status_label(child_record),
                        }
                    )

                selection_df = st.data_editor(
                    pd.DataFrame(selection_rows),
                    hide_index=True,
                    width="stretch",
                    key="stammdaten_child_selection_table",
                    column_config={
                        "Auswahl / Select": st.column_config.CheckboxColumn(
                            "Auswahl / Select"
                        )
                    },
                    disabled=[
                        "Name / Name",
                        "Parent Email",
                        "Group",
                        "Birthdate",
                        "Folder Status",
                    ],
                )
                selected_children: list[dict[str, str]] = []
                selected_child_ids: list[str] = []
                selection_values = selection_df["Auswahl / Select"].tolist()
                for index, selected in enumerate(selection_values):
                    if not bool(selected):
                        continue
                    selected_children.append(children[index])
                    selected_child_ids.append(
                        str(children[index].get("id", "")).strip()
                    )
                st.session_state["stammdaten_selected_child_ids"] = selected_child_ids

                if not selected_children:
                    st.caption(
                        "Bitte mindestens ein Kind in der Tabelle auswählen. / "
                        "Please select at least one child in the table."
                    )
                else:
                    st.markdown(
                        "**Ausgewählte Stammdaten bearbeiten / Edit selected master data**"
                    )
                    child_columns = st.columns(len(selected_children))
                    for index, child_record in enumerate(selected_children):
                        with child_columns[index]:
                            child_name = _display_or_dash(child_record.get("name"))
                            st.markdown(f"### {child_name}")
                            parent_record = (
                                stammdaten_manager.get_parent_by_email(
                                    str(child_record.get("parent_email", "")).strip()
                                )
                                or {}
                            )
                            key_prefix = f"child_{str(child_record.get('id', '')).strip() or index}"
                            _render_child_edit_form(
                                child_record,
                                parent_record,
                                key_prefix,
                                stammdaten_manager,
                            )

                _render_export_backup_section()

            if children:
                with st.expander(
                    "Abholberechtigte / Pickup authorizations", expanded=False
                ):
                    pickup_child = st.selectbox(
                        "Kind für Abholberechtigungen / Child for pickup authorizations",
                        options=children,
                        format_func=lambda child: str(child.get("name", "")),
                        key="pickup_child_select",
                        index=None,
                        placeholder="Kind auswählen / Select child",
                    )
                    if pickup_child is None:
                        st.caption(
                            "Bitte zuerst ein Kind auswählen, um Abholberechtigungen "
                            "anzuzeigen. / Please select a child first to view pickup "
                            "authorizations."
                        )
                        pickup_records = []
                    else:
                        pickup_child_id = str(pickup_child.get("id", "")).strip()
                        pickup_records = (
                            stammdaten_manager.get_pickup_authorizations_by_child_id(
                                pickup_child_id
                            )
                        )

                    if pickup_child is not None and pickup_records:
                        for record in pickup_records:
                            state_label = (
                                "Aktiv / Active"
                                if _normalize_active_flag(record.get("active"))
                                else "Inaktiv / Inactive"
                            )
                            period_parts = []
                            if record.get("valid_from"):
                                period_parts.append(
                                    f"ab / from {record.get('valid_from')}"
                                )
                            if record.get("valid_to"):
                                period_parts.append(
                                    f"bis / until {record.get('valid_to')}"
                                )
                            period_label = (
                                f" ({', '.join(period_parts)})" if period_parts else ""
                            )
                            st.write(
                                f"- **{record.get('name', '—')}** · "
                                f"{record.get('relationship', '—')} · "
                                f"{record.get('phone', '—')} · "
                                f"{state_label}{period_label}"
                            )
                    elif pickup_child is not None:
                        st.caption(
                            "Keine Abholberechtigten hinterlegt. / No pickup authorizations yet."
                        )

                    if pickup_child is not None:
                        with st.form(key="pickup_add_form"):
                            pickup_add_col_left, pickup_add_col_right = st.columns(2)
                            with pickup_add_col_left:
                                pickup_name = st.text_input(
                                    "Name / Name",
                                    key="pickup_add_name",
                                )
                                pickup_relationship = st.text_input(
                                    "Beziehung / Relationship",
                                    key="pickup_add_relationship",
                                )
                                pickup_phone = st.text_input(
                                    "Telefon / Phone",
                                    key="pickup_add_phone",
                                )
                            with pickup_add_col_right:
                                pickup_valid_from = st.date_input(
                                    "Gültig von / Valid from",
                                    value=None,
                                    format="YYYY-MM-DD",
                                    key="pickup_add_valid_from",
                                )
                                pickup_valid_to = st.date_input(
                                    "Gültig bis / Valid to",
                                    value=None,
                                    format="YYYY-MM-DD",
                                    key="pickup_add_valid_to",
                                )
                                pickup_active = st.checkbox(
                                    "Aktiv / Active",
                                    value=True,
                                    key="pickup_add_active",
                                )
                            pickup_add_submitted = st.form_submit_button(
                                "Abholberechtigte hinzufügen / Add pickup authorization"
                            )
                        if pickup_add_submitted:
                            if not pickup_name.strip():
                                st.error("Bitte Name angeben. / Please provide a name.")
                            else:
                                try:
                                    stammdaten_manager.add_pickup_authorization(
                                        pickup_child_id,
                                        {
                                            "name": pickup_name.strip(),
                                            "relationship": pickup_relationship.strip(),
                                            "phone": pickup_phone.strip(),
                                            "valid_from": _optional_date_to_iso(
                                                pickup_valid_from
                                            ),
                                            "valid_to": _optional_date_to_iso(
                                                pickup_valid_to
                                            ),
                                            "active": _active_flag_to_string(
                                                pickup_active
                                            ),
                                        },
                                        created_by=user_email,
                                    )
                                    st.success(
                                        "Abholberechtigung angelegt. / Pickup authorization created."
                                    )
                                    _trigger_rerun()
                                except Exception as exc:
                                    st.error(
                                        "Abholberechtigung konnte nicht gespeichert werden / "
                                        f"Could not save pickup authorization: {exc}"
                                    )

                if pickup_records:
                    selected_pickup_record = st.selectbox(
                        "Abholberechtigte bearbeiten / Edit pickup authorization",
                        options=pickup_records,
                        format_func=lambda record: str(record.get("name", "")),
                        key="pickup_edit_select",
                    )
                    with st.form(key="pickup_edit_form"):
                        pickup_edit_col_left, pickup_edit_col_right = st.columns(2)
                        with pickup_edit_col_left:
                            pickup_edit_name = st.text_input(
                                "Name / Name",
                                value=selected_pickup_record.get("name", ""),
                                key="pickup_edit_name",
                            )
                            pickup_edit_relationship = st.text_input(
                                "Beziehung / Relationship",
                                value=selected_pickup_record.get("relationship", ""),
                                key="pickup_edit_relationship",
                            )
                            pickup_edit_phone = st.text_input(
                                "Telefon / Phone",
                                value=selected_pickup_record.get("phone", ""),
                                key="pickup_edit_phone",
                            )
                        with pickup_edit_col_right:
                            pickup_edit_valid_from = st.date_input(
                                "Gültig von / Valid from",
                                value=_parse_optional_iso_date(
                                    selected_pickup_record.get("valid_from", "")
                                ),
                                format="YYYY-MM-DD",
                                key="pickup_edit_valid_from",
                            )
                            pickup_edit_valid_to = st.date_input(
                                "Gültig bis / Valid to",
                                value=_parse_optional_iso_date(
                                    selected_pickup_record.get("valid_to", "")
                                ),
                                format="YYYY-MM-DD",
                                key="pickup_edit_valid_to",
                            )
                            pickup_edit_active = st.checkbox(
                                "Aktiv / Active",
                                value=_normalize_active_flag(
                                    selected_pickup_record.get("active")
                                ),
                                key="pickup_edit_active",
                            )
                        pickup_edit_submitted = st.form_submit_button(
                            "Abholberechtigung speichern / Save pickup authorization"
                        )
                    if pickup_edit_submitted:
                        pickup_id = str(
                            selected_pickup_record.get("pickup_id", "")
                        ).strip()
                        if not pickup_id:
                            st.error(
                                "Pickup-ID fehlt. / Pickup ID missing in selected record."
                            )
                        elif not pickup_edit_name.strip():
                            st.error("Bitte Name angeben. / Please provide a name.")
                        else:
                            try:
                                stammdaten_manager.update_pickup_authorization(
                                    pickup_id,
                                    {
                                        "name": pickup_edit_name.strip(),
                                        "relationship": pickup_edit_relationship.strip(),
                                        "phone": pickup_edit_phone.strip(),
                                        "valid_from": _optional_date_to_iso(
                                            pickup_edit_valid_from
                                        ),
                                        "valid_to": _optional_date_to_iso(
                                            pickup_edit_valid_to
                                        ),
                                        "active": _active_flag_to_string(
                                            pickup_edit_active
                                        ),
                                    },
                                )
                                st.success(
                                    "Abholberechtigung aktualisiert. / Pickup authorization updated."
                                )
                                _trigger_rerun()
                            except Exception as exc:
                                st.error(
                                    "Aktualisierung fehlgeschlagen / Update failed: "
                                    f"{exc}"
                                )

                    deactivate_label = (
                        "Deaktivieren / Deactivate"
                        if _normalize_active_flag(selected_pickup_record.get("active"))
                        else "Aktivieren / Activate"
                    )
                    if st.button(deactivate_label, key="pickup_toggle_active"):
                        pickup_id = str(
                            selected_pickup_record.get("pickup_id", "")
                        ).strip()
                        if pickup_id:
                            new_active = not _normalize_active_flag(
                                selected_pickup_record.get("active")
                            )
                            stammdaten_manager.update_pickup_authorization(
                                pickup_id,
                                {"active": _active_flag_to_string(new_active)},
                            )
                            st.success("Status aktualisiert. / Pickup status updated.")
                            _trigger_rerun()

            if app_config.storage_mode == "google":
                st.info(
                    "Beim Anlegen eines Kindes wird automatisch ein zugehöriger Drive-Ordner erstellt und verknüpft."
                )
            else:
                st.info(
                    "Beim Anlegen eines Kindes wird lokal ein Prototyp-Ordner erstellt. / "
                    "A local prototype folder is created automatically."
                )

            if children:
                with st.expander("Medikationen", expanded=False):
                    medication_child = st.selectbox(
                        "Kind auswählen / Select child",
                        options=children,
                        format_func=lambda item: str(item.get("name", "")),
                        key="admin_medication_child",
                    )
                    medication_child_id = str(medication_child.get("id", "")).strip()

                    with st.form("admin_medication_form"):
                        med_datetime = st.text_input(
                            "Zeitpunkt (ISO) / Date-time (ISO)",
                            value=pd.Timestamp.now(tz="Europe/Berlin").strftime(
                                "%Y-%m-%dT%H:%M"
                            ),
                            help="Beispiel / Example: 2026-01-15T08:30",
                        )
                        med_name = st.text_input("Medikament / Medication")
                        dose = st.text_input("Dosis / Dose")
                        given_by = st.text_input("Verabreicht von / Given by")
                        notes = st.text_area("Notizen / Notes", height=90)
                        consent_doc_file_id = st.text_input(
                            "Consent-Dokument File-ID (optional) / Consent document file ID (optional)"
                        )
                        add_medication_submitted = st.form_submit_button(
                            "Eintrag speichern / Save entry"
                        )

                    if add_medication_submitted:
                        if (
                            not med_name.strip()
                            or not dose.strip()
                            or not given_by.strip()
                        ):
                            st.error(
                                "Bitte Medikament, Dosis und Verabreicht-von ausfüllen. / "
                                "Please fill medication, dose, and given-by."
                            )
                        else:
                            try:
                                stammdaten_manager.add_medication(
                                    medication_child_id,
                                    {
                                        "date_time": med_datetime.strip(),
                                        "med_name": med_name.strip(),
                                        "dose": dose.strip(),
                                        "given_by": given_by.strip(),
                                        "notes": notes.strip(),
                                        "consent_doc_file_id": consent_doc_file_id.strip(),
                                    },
                                    created_by=user_email,
                                )
                                st.success("Eintrag gespeichert. / Entry saved.")
                                _trigger_rerun()
                            except Exception as exc:
                                st.error(
                                    "Eintrag konnte nicht gespeichert werden. / Could not save entry."
                                )
                                st.info(str(exc))

                    medication_records = stammdaten_manager.get_medications_by_child_id(
                        medication_child_id
                    )
                    st.markdown("**Einträge / Entries**")
                    if medication_records:
                        for record in medication_records:
                            consent_id = str(
                                record.get("consent_doc_file_id", "")
                            ).strip()
                            st.write(
                                f"- **{record.get('date_time', '—')}** · "
                                f"{record.get('med_name', '—')} · "
                                f"{record.get('dose', '—')} · "
                                f"{record.get('given_by', '—')}"
                            )
                            if record.get("notes"):
                                st.caption(f"Notiz / Note: {record.get('notes')}")
                            if consent_id:
                                st.caption(f"Consent Doc File ID: `{consent_id}`")
                            else:
                                st.warning(
                                    "Hinweis: Kein Consent-Dokument verknüpft (optional). / "
                                    "Note: No consent document linked (optional)."
                                )
                    else:
                        st.caption("Noch keine Einträge vorhanden. / No entries yet.")

        # ---- Admin: Dokumente ----
        elif admin_view == "Dokumente":
            st.subheader("Dokumente generieren und verwalten")
            children = stammdaten_manager.get_children()
            if not children:
                st.warning(
                    "Keine Kinder vorhanden. Bitte zuerst Stammdaten anlegen. / "
                    "No children available. Please add master data first."
                )
            else:
                # Formular für Dokumentenerstellung
                sel_child = st.selectbox(
                    "Für welches Kind soll ein Dokument erstellt werden? / "
                    "Select child for document generation",
                    options=children,
                    format_func=lambda child: str(child.get("name", "")),
                )
                doc_notes = st.text_area(
                    "Stichpunkte oder Notizen für das Dokument / Notes for the report"
                )
                save_to_drive = st.checkbox(
                    "Dokument im Drive-Ordner des Kindes speichern? / "
                    "Save document in child's Drive folder?"
                )
                if st.button("Bericht erstellen / Create report"):
                    with st.spinner("Generiere Dokument mit OpenAI..."):
                        try:
                            doc_bytes, file_name = doc_agent.generate_document(
                                sel_child, doc_notes
                            )
                            st.success("Dokument erstellt: " + file_name)
                            with st.expander(
                                "Vorschau des neuen Dokuments / Preview new document",
                                expanded=True,
                            ):
                                st.markdown(_extract_docx_preview_text(doc_bytes))
                            # Download-Button anzeigen
                            st.download_button(
                                "📄 Dokument herunterladen",
                                data=doc_bytes,
                                file_name=file_name,
                            )
                            # Optional: in Drive speichern
                            if save_to_drive:
                                folder_id = sel_child.get("folder_id")
                                if folder_id:
                                    drive_agent.upload_file(
                                        file_name,
                                        doc_bytes,
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        folder_id,
                                    )
                                    st.info(
                                        "Dokument wurde im Drive-Ordner gespeichert."
                                    )
                                else:
                                    st.error(
                                        "Kein Drive-Ordner für dieses Kind vorhanden."
                                    )
                        except DocumentGenerationError as e:
                            st.error(
                                "Dokument konnte nicht erstellt werden. Bitte Hinweise prüfen und erneut versuchen."
                            )
                            st.error(
                                "Document could not be generated. Please review the message and retry."
                            )
                            st.info(str(e))
                        except Exception as e:
                            st.error(f"Fehler bei der Dokumentenerstellung: {e}")

                st.divider()
                st.write("**Vorlagen aus Stammdaten / Templates from master data**")
                contract_col, invoice_col = st.columns(2)

                with contract_col:
                    st.write("Betreuungsvertrag / Childcare contract")
                    if st.button(
                        "Betreuungsvertrag generieren / Generate childcare contract"
                    ):
                        try:
                            contract_bytes, contract_filename = (
                                doc_agent.generate_care_contract(sel_child)
                            )
                            st.success(
                                "Betreuungsvertrag erstellt. / Childcare contract generated."
                            )
                            with st.expander(
                                "Vorschau Betreuungsvertrag / Childcare contract preview",
                                expanded=True,
                            ):
                                st.markdown(_extract_docx_preview_text(contract_bytes))
                            st.download_button(
                                "📄 Vertrag herunterladen / Download contract",
                                data=contract_bytes,
                                file_name=contract_filename,
                                key=f"contract_download_{sel_child.get('id', 'child')}",
                            )
                        except Exception as exc:
                            st.error(
                                "Vertrag konnte nicht erstellt werden. / "
                                "Contract could not be generated."
                            )
                            st.info(str(exc))

                with invoice_col:
                    st.write("Lebensmittelpauschale abrechnen / Food allowance invoice")
                    period_start = st.date_input(
                        "Startdatum / Start date",
                        value=date.today().replace(day=1),
                        key="food_invoice_start",
                    )
                    period_end = st.date_input(
                        "Enddatum / End date",
                        value=date.today(),
                        key="food_invoice_end",
                    )
                    monthly_amount = st.number_input(
                        "Monatliche Pauschale (€) / Monthly allowance (€)",
                        min_value=0.0,
                        value=120.0,
                        step=5.0,
                        key="food_invoice_monthly_amount",
                    )
                    if st.button(
                        "Abrechnung erstellen / Generate invoice",
                        key="create_food_invoice_button",
                    ):
                        try:
                            invoice_bytes, invoice_filename = (
                                doc_agent.generate_food_allowance_invoice(
                                    child_data=sel_child,
                                    period_start=period_start,
                                    period_end=period_end,
                                    monthly_amount_eur=float(monthly_amount),
                                )
                            )
                            st.success(
                                "Abrechnung erstellt. / Food allowance invoice generated."
                            )
                            with st.expander(
                                "Vorschau Abrechnung / Invoice preview",
                                expanded=True,
                            ):
                                st.markdown(_extract_docx_preview_text(invoice_bytes))
                            st.download_button(
                                "📄 Abrechnung herunterladen / Download invoice",
                                data=invoice_bytes,
                                file_name=invoice_filename,
                                key=f"invoice_download_{sel_child.get('id', 'child')}",
                            )
                        except DocumentGenerationError as exc:
                            st.error(
                                "Abrechnung konnte nicht erstellt werden. / "
                                "Invoice could not be generated."
                            )
                            st.info(str(exc))
                        except Exception as exc:
                            st.error(f"Fehler bei der Abrechnungserstellung: {exc}")

                # Optionale Liste vorhandener Dokumente im Drive
                if sel_child.get("folder_id"):
                    docs_list = drive_agent.list_files(
                        sel_child["folder_id"],
                        mime_type_filter="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    if docs_list:
                        st.write("**Bereits gespeicherte Dokumente für dieses Kind:**")
                        for doc in docs_list:
                            file_name = doc.get("name")
                            file_id = doc.get("id")
                            doc_bytes = drive_agent.download_file(file_id)
                            preview_title = (
                                "Vorschau gespeichertes Dokument"
                                " / Preview saved document"
                            )
                            with st.expander(f"📄 {file_name}"):
                                st.caption(preview_title)
                                st.markdown(_extract_docx_preview_text(doc_bytes))
                                st.download_button(
                                    "Download / Herunterladen",
                                    data=doc_bytes,
                                    file_name=file_name,
                                    key=file_id,
                                )
                    else:
                        st.write(
                            "(Keine gespeicherten Dokumente vorhanden. / "
                            "No saved documents available.)"
                        )

        # ---- Admin: Verträge ----
        elif admin_view == "Verträge":
            st.subheader("Vertragsablage / Contract storage")
            if app_config.storage_mode != "google" or app_config.google is None:
                st.info(
                    "Die Vertragsablage ist nur im Google-Modus verfügbar. / "
                    "Contract storage is only available in Google mode."
                )
            else:
                contracts_folder_id = app_config.google.drive_contracts_folder_id
                with st.form("contracts_upload_form", border=True):
                    contract_file = st.file_uploader(
                        "Vertrag hochladen (PDF/DOCX) / Upload contract (PDF/DOCX)",
                        type=["pdf", "docx"],
                        key="contracts_uploader",
                    )
                    contract_upload_submitted = st.form_submit_button(
                        "In Drive speichern / Save to Drive"
                    )

                if contract_upload_submitted:
                    if contract_file is None:
                        st.warning(
                            "Bitte zuerst eine PDF- oder DOCX-Datei auswählen. / "
                            "Please select a PDF or DOCX file first."
                        )
                    else:
                        try:
                            file_id = drive_agent.upload_file(
                                name=contract_file.name,
                                content_bytes=contract_file.getvalue(),
                                mime_type=contract_file.type
                                or "application/octet-stream",
                                parent_folder_id=contracts_folder_id,
                            )
                            st.success(
                                "Datei in Google Drive gespeichert. / "
                                f"File saved to Google Drive (ID: {file_id})."
                            )
                        except DriveServiceError as exc:
                            st.error(
                                "Upload fehlgeschlagen. Prüfen Sie die Ordnerfreigabe für "
                                "den Service-Account (403/404). / Upload failed. "
                                "Please verify folder sharing with the service account "
                                "(403/404)."
                            )
                            st.info(str(exc))

                st.write("**Vorhandene Vertragsdateien / Existing contract files**")
                try:
                    contract_files = drive_agent.list_files(contracts_folder_id)
                    if contract_files:
                        for file_meta in contract_files:
                            st.markdown(
                                f"- **{file_meta.get('name', '-')}** "
                                f"`{file_meta.get('mimeType', '-')}` · "
                                f"{file_meta.get('modifiedTime', '-')}"
                            )
                    else:
                        st.caption(
                            "Noch keine Dateien vorhanden. / No files available yet."
                        )
                except DriveServiceError as exc:
                    st.error(
                        "Dateiliste konnte nicht geladen werden. Stellen Sie sicher, "
                        "dass der Zielordner mit dem Service-Account geteilt ist. / "
                        "Could not load file list. Ensure the folder is shared with "
                        "the service account."
                    )
                    st.info(str(exc))

        # ---- Admin: Fotos ----
        elif admin_view == "Fotos":
            st.subheader("Fotos & Videos verwalten / Manage photos & videos")
            children = stammdaten_manager.get_children()
            if not children:
                st.warning(
                    "Bitte legen Sie zuerst Kinder-Stammdaten an. / "
                    "Please create child records first."
                )
            else:
                render_media_page(
                    MediaPageContext(
                        app_config=app_config,
                        user_email=user_email,
                        children=children,
                        stammdaten_manager=stammdaten_manager,
                        drive_agent=drive_agent,
                        ensure_child_photo_folder=ensure_child_photo_folder,
                        trigger_rerun=_trigger_rerun,
                    )
                )

        # ---- Admin: Kalender ----
        elif admin_view == "Kalender":
            st.subheader("Kalenderansicht / Calendar view")
            components.html(GOOGLE_CALENDAR_EMBED_HTML, height=320)

            st.subheader("Neuer Termin / New event")
            title = st.text_input("Titel / Title")
            event_date = st.date_input("Datum / Date")
            event_time = st.time_input("Uhrzeit (Start) / Start time")
            description = st.text_area("Beschreibung / Description")
            if st.button("Termin hinzufügen / Add event"):
                if not title:
                    st.error("Bitte Titel eingeben. / Please enter a title.")
                else:
                    try:
                        add_event(
                            title=title,
                            event_date=event_date,
                            event_time=event_time,
                            description=description,
                        )
                        st.success("Termin wurde hinzugefügt. / Event created.")
                    except CalendarServiceError as exc:
                        st.error(f"Fehler beim Speichern / Failed to save event: {exc}")

            st.write("**Bevorstehende Termine / Upcoming events**")
            try:
                events = list_events(max_results=10)
            except CalendarServiceError as exc:
                events = []
                st.error(f"Fehler beim Laden / Failed to load events: {exc}")
            if events:
                for ev in events:
                    st.write(f"- {ev['start']} · **{ev['summary']}**")
                    if ev["description"]:
                        st.caption(ev["description"])
            else:
                st.write("Keine anstehenden Termine vorhanden. / No upcoming events.")

        elif admin_view == "System / Healthchecks":
            st.subheader("System / Healthchecks")
            st.caption(
                "Prüfen Sie die Integrationen zu Google Drive, Kalender und Sheets. / "
                "Check integrations for Google Drive, Calendar and Sheets."
            )
            if healthcheck_requested:
                with st.spinner(
                    "Prüfe Drive, Kalender & Sheets... / "
                    "Checking drive, calendar & sheets..."
                ):
                    check_results = _run_google_connection_check()
                for check_title, ok, message in check_results:
                    if ok:
                        st.success(f"{check_title}: {message}")
                    else:
                        st.error(f"{check_title}: {message}")
            else:
                st.info(
                    "Nutzen Sie den Button in der Sidebar, um den Check zu starten. / "
                    "Use the sidebar button to run the check."
                )

    else:
        # ---- Parent/Eltern View ----
        child = stammdaten_manager.get_child_by_parent(user_email)
        st.session_state.child = child
        if menu == "child":
            with st.container(border=True):
                st.subheader("Mein Kind - Übersicht")
                if child:
                    parent_record = (
                        stammdaten_manager.get_parent_by_email(
                            str(child.get("parent_email", "")).strip()
                        )
                        or {}
                    )
                    st.markdown("**Kinderdetails / Child details**")
                    st.write(f"**Name:** {_display_or_dash(child.get('name'))}")
                    st.write(
                        "**Geburtsdatum / Birthdate:** "
                        f"{_display_or_dash(child.get('birthdate'))}"
                    )
                    if child.get("start_date"):
                        st.write(
                            f"**Startdatum / Start date:** {child.get('start_date')}"
                        )
                    st.write(
                        f"**Gruppe / Group:** {_display_or_dash(child.get('group'))}"
                    )
                    if child.get("primary_caregiver"):
                        st.write(
                            f"**Bezugserzieher:in / Primary caregiver:** {child.get('primary_caregiver')}"
                        )
                    if child.get("allergies"):
                        st.write(f"**Allergien / Allergies:** {child.get('allergies')}")
                    if child.get("notes_parent_visible"):
                        st.info(
                            f"Hinweise / Notes:\n\n{child.get('notes_parent_visible')}"
                        )

                    child_pickup_records = (
                        stammdaten_manager.get_pickup_authorizations_by_child_id(
                            str(child.get("id", "")).strip(),
                            active_only=True,
                        )
                    )
                    st.markdown("**Wer ist eingetragen? / Who is listed?**")
                    if child_pickup_records:
                        for record in child_pickup_records:
                            period_parts = []
                            if record.get("valid_from"):
                                period_parts.append(
                                    f"ab / from {record.get('valid_from')}"
                                )
                            if record.get("valid_to"):
                                period_parts.append(
                                    f"bis / until {record.get('valid_to')}"
                                )
                            period_label = (
                                f" ({', '.join(period_parts)})" if period_parts else ""
                            )
                            st.write(
                                f"- **{record.get('name', '—')}** · "
                                f"{record.get('relationship', '—')} · "
                                f"{record.get('phone', '—')}{period_label}"
                            )
                    else:
                        st.caption(
                            "Keine aktiven Abholberechtigten hinterlegt. / "
                            "No active pickup authorizations listed."
                        )

                    st.markdown("**Elterninformationen / Parent information**")
                    emergency_name = str(
                        parent_record.get("emergency_contact_name", "")
                    ).strip()
                    emergency_phone = str(
                        parent_record.get("emergency_contact_phone", "")
                    ).strip()
                    emergency_parts = [
                        part for part in [emergency_name, emergency_phone] if part
                    ]
                    st.write(
                        "**Notfallkontakt / Emergency contact:** "
                        f"{' – '.join(emergency_parts) if emergency_parts else '-'}"
                    )
                    st.write(
                        "**Bevorzugte Sprache / Preferred language:** "
                        f"{_language_label(str(parent_record.get('preferred_language', '')))}"
                    )
                    st.write(
                        "**Benachrichtigungen / Notifications:** "
                        f"{'Ja / Yes' if _parse_opt_in_flag(parent_record.get('notifications_opt_in')) else 'Nein / No'}"
                    )
                else:
                    st.write("Keine Kinderdaten gefunden. / No child data found.")
        elif menu == "info":
            st.subheader("Infos / Information")
            language = st.radio(
                "Sprache / Language",
                options=["de", "en"],
                horizontal=True,
                format_func=lambda value: "Deutsch" if value == "de" else "English",
                key="parent_infos_language",
            )
            try:
                pages = [
                    page
                    for page in content_repo.list_pages()
                    if page.published and page.audience in {"parent", "both"}
                ]
            except ContentRepositoryError as exc:
                pages = []
                st.error(
                    "Infos konnten nicht geladen werden. / Could not load info pages."
                )
                st.info(str(exc))

            if not pages:
                st.info(
                    "Derzeit sind keine veröffentlichten Infos verfügbar. / "
                    "No published info pages are currently available."
                )
            else:
                selected_slug = st.selectbox(
                    "Seite auswählen / Select page",
                    options=[page.slug for page in pages],
                )
                selected_page = next(
                    page for page in pages if page.slug == selected_slug
                )
                st.markdown(f"### {_page_title(selected_page, language)}")
                page_content = _page_body(selected_page, language)
                if page_content:
                    st.markdown(page_content)
                else:
                    st.caption("Kein Inhalt vorhanden. / No content available.")

        elif menu == "documents":
            st.subheader("Dokumente Ihres Kindes")
            if child and child.get("folder_id"):
                docs_list = drive_agent.list_files(
                    child["folder_id"],
                    mime_type_filter="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                if docs_list:
                    for doc in docs_list:
                        file_name = doc.get("name")
                        file_id = doc.get("id")
                        st.markdown(f"**{file_name}** ")
                        st.download_button(
                            "Herunterladen",
                            data=drive_agent.download_file(file_id),
                            file_name=file_name,
                            key=file_id,
                        )
                else:
                    st.write("Keine Dokumente vorhanden.")
            else:
                st.write("Keine Dokumente verfügbar.")
        elif menu == "photos":
            st.subheader("Fotos / Photos")
            if child and child.get("id"):
                try:
                    if app_config.storage_mode == "google":
                        photo_folder_id = str(
                            child.get("photo_folder_id") or child.get("folder_id") or ""
                        ).strip()
                    else:
                        photo_folder_id = str(
                            child.get("photo_folder_id") or child.get("folder_id") or ""
                        )
                    photos = (
                        drive_agent.list_files(
                            photo_folder_id, mime_type_filter="image/"
                        )
                        if photo_folder_id
                        else []
                    )
                    published_photos: list[dict[str, str]] = []
                    for photo in photos:
                        file_id = str(photo.get("id", "")).strip()
                        if not file_id:
                            continue
                        meta = stammdaten_manager.get_photo_meta_by_file_id(file_id)
                        status = _normalize_photo_status(
                            None if meta is None else meta.get("status")
                        )
                        if status == "published":
                            published_photos.append(photo)
                    photos = published_photos
                except Exception as exc:
                    photos = []
                    st.error(
                        "Fotos konnten nicht geladen werden. / Could not load photos."
                    )
                    st.info(str(exc))

                if child:
                    current_download_consent = (
                        str(child.get("download_consent", "pixelated")).strip().lower()
                    )
                    consent_options = ["pixelated", "unpixelated", "denied"]
                    if current_download_consent not in consent_options:
                        current_download_consent = "pixelated"
                    selected_download_consent = st.selectbox(
                        "Foto-Download Consent / Photo download consent",
                        options=consent_options,
                        index=consent_options.index(current_download_consent),
                        format_func=lambda mode: (
                            "Downloads verpixelt / Downloads pixelated"
                            if mode == "pixelated"
                            else (
                                "Downloads unverpixelt / Downloads unpixelated"
                                if mode == "unpixelated"
                                else "Download verweigert / Download denied"
                            )
                        ),
                        key="parent_download_consent_select",
                    )
                    if selected_download_consent != current_download_consent:
                        try:
                            stammdaten_manager.update_child(
                                str(child.get("id", "")),
                                {"download_consent": selected_download_consent},
                            )
                            st.success("Consent aktualisiert. / Consent updated.")
                            _trigger_rerun()
                        except Exception as exc:
                            st.error(
                                "Consent konnte nicht gespeichert werden. / Could not save consent."
                            )
                            st.info(str(exc))
                    st.caption(
                        "In der App bleibt die Vorschau unverändert. Der Consent betrifft nur den Download. / "
                        "In-app preview remains unchanged. Consent affects download only."
                    )
                if photos:
                    active_consent_mode = (
                        str((child or {}).get("download_consent", "pixelated"))
                        .strip()
                        .lower()
                    )
                    for photo in photos:
                        file_name = str(photo.get("name", "photo"))
                        file_id = str(photo.get("id", ""))
                        img_bytes = drive_agent.download_file(file_id)
                        st.image(
                            img_bytes,
                            caption=f"{file_name} · Vorschau / Preview",
                            width="stretch",
                        )
                        if active_consent_mode == "denied":
                            st.caption("Download nicht erlaubt / Download not allowed.")
                        else:
                            download_bytes = _get_photo_download_bytes(
                                file_id=file_id,
                                consent_mode=active_consent_mode,
                            )
                            st.download_button(
                                "Foto herunterladen / Download photo",
                                data=download_bytes,
                                file_name=file_name,
                                key=f"download_photo_{file_id}_{active_consent_mode}",
                            )
                else:
                    st.write(
                        "Keine veröffentlichten Fotos vorhanden. / No published photos available."
                    )
            else:
                st.write("Keine Fotos verfügbar. / No photos available.")
        elif menu == "medication":
            st.subheader("Medikamentengabe / Medication log")
            if child and child.get("id"):
                child_id = str(child.get("id", "")).strip()
                medication_records = stammdaten_manager.get_medications_by_child_id(
                    child_id
                )
                if medication_records:
                    for record in medication_records:
                        st.write(
                            f"- **{record.get('date_time', '—')}** · "
                            f"{record.get('med_name', '—')} · "
                            f"{record.get('dose', '—')} · "
                            f"{record.get('given_by', '—')}"
                        )
                        if record.get("notes"):
                            st.caption(f"Notiz / Note: {record.get('notes')}")
                        if not str(record.get("consent_doc_file_id", "")).strip():
                            st.caption(
                                "Hinweis: Kein Consent-Dokument hinterlegt. / "
                                "Note: No consent document linked."
                            )
                else:
                    st.write(
                        "Keine Medikamenten-Einträge vorhanden. / "
                        "No medication entries available."
                    )
            else:
                st.write("Keine Kinderdaten gefunden. / No child data found.")

        elif menu == "appointments":
            st.subheader("Termine / Events")
            components.html(GOOGLE_CALENDAR_EMBED_HTML, height=320)
            try:
                events = list_events(max_results=10)
            except CalendarServiceError as exc:
                events = []
                st.error(f"Fehler beim Laden / Failed to load events: {exc}")
            if events:
                for ev in events:
                    st.write(f"- {ev['start']} · **{ev['summary']}**")
                    if ev["description"]:
                        st.caption(ev["description"])
            else:
                st.write(
                    "Zurzeit sind keine Termine eingetragen. / No events available right now."
                )
